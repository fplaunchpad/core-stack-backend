import os
import re
import socket
import ssl
from collections import defaultdict
from datetime import date
from io import BytesIO

import geopandas as gpd
from django.core.mail import EmailMessage
from django.core.mail.backends.smtp import EmailBackend
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from shapely.geometry import Point

from dpr.mapping import populate_maintenance_from_waterbody
from dpr.utils import ensure_str, get_waterbody_repair_activities, transform_name
from moderation.views import sync_odk_to_csdb
from nrm_app.settings import (
    DEBUG,
    EMAIL_HOST,
    EMAIL_HOST_PASSWORD,
    EMAIL_HOST_USER,
    EMAIL_PORT,
    EMAIL_TIMEOUT,
    EMAIL_USE_SSL,
    GEOSERVER_URL,
    TMP_LOCATION,
)
from plans.models import PlanApp
from utilities.logger import setup_logger

from .models import (
    ODK_settlement,
    ODK_waterbody,
    ODK_well,
)
from .services import (
    get_crops_data,
    get_livelihood_data,
    get_livestock_data,
    get_maintenance_data,
    get_nrm_works_data,
)
from .utils import (
    format_text,
    format_text_demands,
    get_vector_layer_geoserver,
    sort_key,
    to_utf8,
    transform_name,
)

logger = setup_logger(__name__)


def get_plan_details(plan_id):
    try:
        return PlanApp.objects.get(id=plan_id)
    except PlanApp.DoesNotExist:
        return None


def create_dpr_document(plan):
    logger.info("Generating DPR for plan ID: %s", plan.id)

    doc = initialize_document()  # doc init

    logger.info("Database sync started")  # db sync
    sync_odk_to_csdb()
    logger.info("Database sync complete")

    logger.info("Plan Details")
    logger.info(plan)
    logger.info(transform_name(str(plan.district_soi.district_name)))
    logger.info(transform_name(str(plan.tehsil_soi.tehsil_name)))

    total_settlements = get_settlement_count_for_plan(plan.id)
    logger.info(f"Total settlements found in the plan: {total_settlements}")

    mws_fortnight = get_vector_layer_geoserver(
        geoserver_url=GEOSERVER_URL,
        workspace="mws_layers",
        layer_name="deltaG_fortnight_"
        + transform_name(str(plan.district_soi.district_name))
        + "_"
        + transform_name(str(plan.tehsil_soi.tehsil_name)),
    )
    logger.info(f"MWS Fortnight layer fetched successfully")

    # MARK: - Sections
    add_section_a(doc, plan)
    add_section_separator(doc)
    logger.info("Section A completed")

    settlement_mws_ids, mws_gdf = add_section_b(
        doc, plan, total_settlements, mws_fortnight
    )
    add_section_separator(doc)
    logger.info("Section B completed")

    add_section_c(doc, plan)
    add_section_separator(doc)
    logger.info("Section C completed")

    add_section_d(doc, plan, settlement_mws_ids, mws_gdf)
    add_section_separator(doc)
    logger.info("Section D completed")

    add_section_e(doc, plan)
    add_section_separator(doc)
    logger.info("Section E completed")

    add_section_f(doc, plan, mws_fortnight)
    add_section_separator(doc)
    logger.info("Section F completed")

    add_section_g(doc, plan, mws_fortnight)
    add_section_separator(doc)
    logger.info("Section G completed")

    # MARK: local save /var/www/tmp/dpr/
    if DEBUG:
        file_path = TMP_LOCATION + "dpr/"

        if not os.path.exists(file_path):
            os.makedirs(file_path)
        doc.save(file_path + plan.plan + ".docx")
    return doc


def get_data_for_settlement(planid):
    return (
        ODK_settlement.objects.filter(plan_id=planid)
        .exclude(status_re="rejected")
        .exclude(is_deleted=True)
    )


def get_settlement_count_for_plan(planid):
    return (
        ODK_settlement.objects.filter(plan_id=planid)
        .exclude(status_re="rejected")
        .exclude(is_deleted=True)
        .count()
    )


def get_settlement_coordinates_for_plan(planid):
    settlements = (
        ODK_settlement.objects.filter(plan_id=planid)
        .exclude(status_re="rejected")
        .exclude(is_deleted=True)
        .values("settlement_name", "latitude", "longitude")
    )
    return [
        (
            settlement["settlement_name"],
            settlement["latitude"],
            settlement["longitude"],
        )
        for settlement in settlements
    ]


def get_mws_uid_for_settlement_gdf(mws_gdf, lat, lon):
    settlement_point = Point(lon, lat)
    intersecting_mws = mws_gdf[mws_gdf.intersects(settlement_point)]

    if not intersecting_mws.empty:
        mws_uid = intersecting_mws.iloc[0]["uid"]
        return mws_uid
    else:
        return None


def add_section_separator(doc):
    """Add a centered *** separator with padding and larger font size."""
    doc.add_paragraph()
    para = doc.add_paragraph("***")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in para.runs:
        run.font.size = Pt(18)
    doc.add_paragraph()


def initialize_document():
    doc = Document()
    heading = doc.add_heading("Detailed Project Report", 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        date.today().strftime("%B %d, %Y")
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return doc


def get_mws_ids_for_report(plan):
    mws_ids = set()
    settlement_coordinates = get_settlement_coordinates_for_plan(plan.id)

    mws_fortnight = get_vector_layer_geoserver(
        geoserver_url=GEOSERVER_URL,
        workspace="mws_layers",
        layer_name="deltaG_fortnight_"
        + transform_name(str(plan.district_soi.district_name))
        + "_"
        + transform_name(str(plan.tehsil_soi.tehsil_name)),
    )
    mws_gdf = gpd.GeoDataFrame.from_features(mws_fortnight["features"])

    for settlement_name, latitude, longitude in settlement_coordinates:
        mws_uid = get_mws_uid_for_settlement_gdf(mws_gdf, latitude, longitude)
        if mws_uid:
            mws_ids.add(mws_uid)

    return sorted(mws_ids)


############################## Sections #######################################


# MARK: -  A
def add_section_a(doc, plan):
    """
    Brief about team details.
    """
    doc.add_heading("Section A: Team Details", level=1)
    create_table_team_details(doc, plan)


def create_table_team_details(doc, plan):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    row1_cells = table.rows[0].cells
    row1_cells[0].text = "Organization"
    row1_cells[1].text = to_utf8(plan.organization.name) if plan.organization else "NA"

    row2_cells = table.rows[1].cells
    row2_cells[0].text = "Project"
    row2_cells[1].text = to_utf8(plan.project.name) if plan.project else "NA"

    row3_cells = table.rows[2].cells
    row3_cells[0].text = "Plan"
    row3_cells[1].text = to_utf8(plan.plan)

    row4_cells = table.rows[3].cells
    row4_cells[0].text = "Facilitator"
    row4_cells[1].text = to_utf8(plan.facilitator_name)

    row5_cells = table.rows[4].cells
    row5_cells[0].text = "Process involved in the preparation of DPR PRA"
    row5_cells[1].text = "PRA, Gram Sabha, Transect Walk, GIS Mapping"

    for row in table.rows:
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


# MARK: - Section B
def add_section_b(doc, plan, total_settlements, mws_fortnight):
    """
    Briefs about the village
    """
    doc.add_heading("Section B: Brief of Village")

    mws_gdf = gpd.GeoDataFrame.from_features(mws_fortnight["features"])

    settlement_mws_ids = []
    settlement_coordinates = get_settlement_coordinates_for_plan(plan.id)

    for settlement_name, latitude, longitude in settlement_coordinates:
        mws_uid = get_mws_uid_for_settlement_gdf(mws_gdf, latitude, longitude)
        if mws_uid:
            settlement_mws_ids.append((settlement_name, mws_uid))

    intersecting_mws_ids = "; ".join(
        [f"{name}: {mws_id}" for name, mws_id in settlement_mws_ids]
    )
    create_table_village_brief(
        intersecting_mws_ids, doc, plan, total_settlements, mws_gdf
    )

    return settlement_mws_ids, mws_gdf


def create_table_village_brief(
    intersecting_mws_ids, doc, plan, total_settlements, mws_gdf
):
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"

    # Calculate the centroid of the intersecting MWS
    if intersecting_mws_ids:
        intersecting_mws = mws_gdf[
            mws_gdf["uid"].isin(
                [mws_id.split(": ")[1] for mws_id in intersecting_mws_ids.split("; ")]
            )
        ]
        if not intersecting_mws.empty:
            centroid = intersecting_mws.geometry.unary_union.centroid
        else:
            centroid = None
    else:
        centroid = None

    headers_data = [
        ("Name of the Village", to_utf8(plan.village_name)),
        ("Name of the Gram Panchayat", to_utf8(plan.gram_panchayat)),
        ("Tehsil", to_utf8(plan.tehsil_soi.tehsil_name)),
        ("District", to_utf8(plan.district_soi.district_name)),
        ("State", to_utf8(plan.state_soi.state_name)),
        ("Number of Settlements in the Village", str(total_settlements)),
        ("Intersecting Micro Watershed IDs", None),  # Will be handled separately
        (
            "Latitude and Longitude of the Village",
            f"{centroid.y:.8f}, {centroid.x:.8f}" if centroid else "Not available",
        ),
    ]

    for i, (key, value) in enumerate(headers_data):
        row_cells = table.rows[i].cells
        row_cells[0].text = key

        if key == "Intersecting Micro Watershed IDs":
            settlement_mws_pairs = []
            if intersecting_mws_ids:
                pairs = intersecting_mws_ids.split("; ")
                for pair in pairs:
                    if ": " in pair:
                        settlement, mws_id = pair.split(": ", 1)
                        settlement_mws_pairs.append((settlement, mws_id))

            if settlement_mws_pairs:
                nested_table = row_cells[1].add_table(
                    rows=len(settlement_mws_pairs) + 1, cols=2
                )
                nested_table.style = "Table Grid"

                header_cells = nested_table.rows[0].cells
                header_cells[0].text = "Settlement"
                header_cells[1].text = "MWS ID"

                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                for idx, (settlement, mws_id) in enumerate(settlement_mws_pairs, 1):
                    data_cells = nested_table.rows[idx].cells
                    data_cells[0].text = to_utf8(settlement)
                    data_cells[1].text = to_utf8(mws_id)
            else:
                row_cells[1].text = "No intersecting watersheds"
        else:
            row_cells[1].text = to_utf8(value) if value else "NA"

    for row in table.rows:
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


# MARK: - Section C
def add_section_c(doc, plan):
    """
    Adds information on Settlement, Vulnerability, NREGA from the settlement's form

    Livelihood Profile: From Livelihood form and Settlement's form (Livestock Info)
    """
    doc.add_heading("Section C: Social Economic Ecological Profile", level=1)
    para = doc.add_paragraph()
    para.add_run(
        "This section includes information on vulnerabilities, NREGA details, livelihoods, crop and livestock profiles."
    )

    settlement_data = get_data_for_settlement(plan.id)
    create_table_socio_eco(doc, plan, settlement_data)

    doc.add_heading("MGNREGA Info", level=4)
    create_table_mgnrega_info(doc, plan, settlement_data)

    doc.add_heading("Crop Info", level=4)
    create_table_crop_info(doc, plan)

    doc.add_heading("Livestock Info", level=4)
    create_table_livestock(doc, plan)


def create_table_socio_eco(doc, plan, settlement_data):
    headers_socio = [
        "Name of the Settlement",
        "Total Number of Households",
        "Settlement Type",
        "Caste Group",
        "Total Households",
        "Total marginal farmers (<2 acres)",
    ]

    table_socio = doc.add_table(rows=1, cols=len(headers_socio))
    table_socio.style = "Table Grid"

    hdr_cells = table_socio.rows[0].cells
    for i, header in enumerate(headers_socio):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for item in settlement_data:
        row_cells = table_socio.add_row().cells
        row_cells[0].text = to_utf8(item.settlement_name)
        row_cells[1].text = str(item.number_of_households)
        largest_caste = (item.largest_caste or "").lower()
        row_cells[2].text = to_utf8(item.largest_caste)

        if largest_caste == "single caste group":
            row_cells[3].text = to_utf8(item.smallest_caste)
        elif largest_caste == "mixed caste group":
            row_cells[3].text = to_utf8(item.settlement_status)
        else:
            row_cells[3].text = "NA"

        sub_table = row_cells[4].add_table(rows=4, cols=2)
        sub_table.style = "Table Grid"
        sub_table.autofit = False
        sub_table.allow_autofit = False

        count_sc = str(item.data_settlement.get("count_sc", "")) or "NA"
        count_st = str(item.data_settlement.get("count_st", "")) or "NA"
        count_obc = str(item.data_settlement.get("count_obc", "")) or "NA"
        count_general = str(item.data_settlement.get("count_general", "")) or "NA"

        caste_data = [
            ("SC", count_sc),
            ("ST", count_st),
            ("OBC", count_obc),
            ("General", count_general),
        ]

        for i, (label, value) in enumerate(caste_data):
            label_cell = sub_table.rows[i].cells[0]
            value_cell = sub_table.rows[i].cells[1]

            label_cell.text = label
            label_cell.width = Pt(50)
            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            label_para.runs[0].bold = True
            label_para.runs[0].font.size = Pt(9)

            value_cell.text = to_utf8(value)
            value_cell.width = Pt(40)
            value_para = value_cell.paragraphs[0]
            value_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            value_para.runs[0].font.size = Pt(9)

        row_cells[5].text = str(item.farmer_family.get("marginal_farmers", "")) or "NA"


def create_table_mgnrega_info(doc, plan, settlement_data):
    headers_nrega = [
        "Settlement's Name",
        "Total Households - applied \n- have NREGA job cards in previous year",
        "Total work days in previous year",
        "Work demands made in previous year",
        "Were you involved in the village level planning?",
        "Issues",
    ]

    table_nrega = doc.add_table(rows=1, cols=len(headers_nrega))
    table_nrega.style = "Table Grid"
    hdr_cells = table_nrega.rows[0].cells
    for i, header in enumerate(headers_nrega):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for settlement_nrega in settlement_data:
        row_cells = table_nrega.add_row().cells
        row_cells[0].text = to_utf8(settlement_nrega.settlement_name)
        row_cells[1].text = (
            "applied: "
            + (
                "NA"
                if settlement_nrega.nrega_job_applied == 0
                else str(settlement_nrega.nrega_job_applied)
            )
            + "\n"
            + "having: "
            + (
                "NA"
                if settlement_nrega.nrega_job_card == 0
                else str(settlement_nrega.nrega_job_card)
            )
        )
        row_cells[2].text = (
            "NA"
            if settlement_nrega.nrega_work_days == 0
            else str(settlement_nrega.nrega_work_days)
        )
        row_cells[3].text = to_utf8(
            format_text_demands(settlement_nrega.nrega_past_work)
        )
        row_cells[4].text = to_utf8(format_text(settlement_nrega.nrega_demand))
        row_cells[5].text = to_utf8(format_text(settlement_nrega.nrega_issues))


def create_table_crop_info(doc, plan):
    headers_cropping_pattern = [
        "Name of the Settlement",
        "Irrigation Source",
        "Crops grown in Kharif",
        "Kharif acreage (acres)",
        "Crops grown in Rabi",
        "Rabi acreage (acres)",
        "Crops grown in Zaid",
        "Zaid acreage (acres)",
        "Cropping Intensity",
        "Land Classification",
    ]
    table_cropping_pattern = doc.add_table(rows=1, cols=len(headers_cropping_pattern))
    table_cropping_pattern.style = "Table Grid"
    hdr_cells = table_cropping_pattern.rows[0].cells
    for i, header in enumerate(headers_cropping_pattern):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    def _fmt_acres(val):
        return str(val) if val is not None else "NA"

    for crop in get_crops_data(plan.id):
        row_cells = table_cropping_pattern.add_row().cells
        row_cells[0].text = to_utf8(crop["beneficiary_settlement"])
        row_cells[1].text = to_utf8(crop["irrigation_source"])
        row_cells[2].text = to_utf8(format_text(crop["kharif_crops"]))
        row_cells[3].text = _fmt_acres(crop["kharif_acres"])
        row_cells[4].text = to_utf8(format_text(crop["rabi_crops"]))
        row_cells[5].text = _fmt_acres(crop["rabi_acres"])
        row_cells[6].text = to_utf8(format_text(crop["zaid_crops"]))
        row_cells[7].text = _fmt_acres(crop["zaid_acres"])
        row_cells[8].text = to_utf8(crop["cropping_intensity"])
        row_cells[9].text = to_utf8(crop["land_classification"])


def create_table_livestock(doc, plan):
    headers_livelihood = [
        "Name of the Settlement",
        "Goats",
        "Sheep",
        "Cattle",
        "Piggery",
        "Poultry",
    ]
    livestock_table = doc.add_table(rows=1, cols=len(headers_livelihood))
    livestock_table.style = "Table Grid"
    hdr_cells = livestock_table.rows[0].cells
    for i, header in enumerate(headers_livelihood):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for item in get_livestock_data(plan.id):
        row_cells = livestock_table.add_row().cells
        row_cells[0].text = to_utf8(item["settlement_name"])
        for i, lt in enumerate(
            ["goats", "sheep", "cattle", "piggery", "poultry"], start=1
        ):
            row_cells[i].text = to_utf8(item[lt] or "NA")


# MARK: - Section D
def add_section_d(doc, plan, settlement_mws_ids, mws_gdf):
    doc.add_heading(
        "Section D: Wells and Water Structures Details",
        level=1,
    )
    para = doc.add_paragraph()
    para.add_run(
        "This section provides details on wells, water structures and household beneficiaries"
    )

    unique_mws_ids = sorted(set([mws_id for _, mws_id in settlement_mws_ids]))
    create_table_mws(doc, plan, settlement_mws_ids, mws_gdf, unique_mws_ids)

    populate_consolidated_water_structures(doc, plan, unique_mws_ids, mws_gdf)


def populate_consolidated_water_structures(doc, plan, unique_mws_ids, mws_gdf):
    all_wells_with_mws = get_all_wells_with_mws(plan, unique_mws_ids, mws_gdf)
    all_waterbodies_with_mws = get_all_waterbodies_with_mws(
        plan, unique_mws_ids, mws_gdf
    )

    populate_consolidated_well_tables(doc, all_wells_with_mws)
    populate_consolidated_waterbody_tables(doc, all_waterbodies_with_mws)


def create_table_mws(doc, plan, settlement_mws_ids, mws_gdf, unique_mws_ids):
    headers_mws = ["Microwatershed ID", "Latitude and Longitude (Centroid)"]
    table_mws = doc.add_table(rows=len(unique_mws_ids) + 1, cols=len(headers_mws))
    table_mws.style = "Table Grid"

    hdr_cells = table_mws.rows[0].cells
    for i, header in enumerate(headers_mws):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for i, mws_id in enumerate(unique_mws_ids, start=1):
        row_cells = table_mws.rows[i].cells
        row_cells[0].text = to_utf8(mws_id)

        matching_feature = mws_gdf[mws_gdf["uid"] == mws_id]
        if not matching_feature.empty:
            centroid = matching_feature.geometry.centroid.iloc[0]
            row_cells[1].text = f"{centroid.y:.8f}, {centroid.x:.8f}"
        else:
            row_cells[1].text = "NA"


def get_all_wells_with_mws(plan, unique_mws_ids, mws_gdf):
    """Get all wells across all MWS IDs with their corresponding MWS assignment"""
    wells_in_plan = (
        ODK_well.objects.filter(plan_id=plan.id)
        .exclude(status_re="rejected")
        .exclude(is_deleted=True)
    )
    all_wells_with_mws = []

    for well in wells_in_plan:
        well_point = Point(well.longitude, well.latitude)
        well_mws_id = None

        for mws_id, mws_polygon in zip(mws_gdf["uid"], mws_gdf["geometry"]):
            if mws_id in unique_mws_ids and well_point.within(mws_polygon):
                well_mws_id = mws_id
                break

        if well_mws_id:
            all_wells_with_mws.append((well, well_mws_id))
        else:
            all_wells_with_mws.append((well, "N/A"))
            logger.info(
                f"Well at ({well.latitude}, {well.longitude}) does not belong to any MWS"
            )

    return all_wells_with_mws


def get_all_waterbodies_with_mws(plan, unique_mws_ids, mws_gdf):
    """Get all waterbodies across all MWS IDs with their corresponding MWS assignment"""
    waterbodies_in_plan = (
        ODK_waterbody.objects.filter(plan_id=plan.id)
        .exclude(status_re="rejected")
        .exclude(is_deleted=True)
    )
    all_waterbodies_with_mws = []

    for waterbody in waterbodies_in_plan:
        waterbody_point = Point(waterbody.longitude, waterbody.latitude)
        waterbody_mws_id = None

        for mws_id, mws_polygon in zip(mws_gdf["uid"], mws_gdf["geometry"]):
            if mws_id in unique_mws_ids and waterbody_point.within(mws_polygon):
                waterbody_mws_id = mws_id
                break

        if waterbody_mws_id:
            all_waterbodies_with_mws.append((waterbody, waterbody_mws_id))
        else:
            all_waterbodies_with_mws.append((waterbody, "N/A"))
            logger.info(
                f"Waterbody at ({waterbody.latitude}, {waterbody.longitude}) does not belong to any MWS"
            )

    return all_waterbodies_with_mws


def populate_consolidated_well_tables(doc, all_wells_with_mws):
    """Create consolidated well summary and detail tables for all MWS IDs"""
    if not all_wells_with_mws:
        doc.add_heading("Well Information", level=3)
        doc.add_paragraph("No wells found in any of the MWS areas.")
        return

    wells_count = defaultdict(int)
    households_count = defaultdict(int)

    for well, mws_id in all_wells_with_mws:
        wells_count[well.beneficiary_settlement] += 1
        households_count[well.beneficiary_settlement] += int(well.households_benefitted)

    wells_info = [
        (settlement, wells_count[settlement], households_count[settlement])
        for settlement in sorted(wells_count.keys(), key=sort_key)
    ]

    doc.add_heading("Well Summary Information", level=3)
    headers_well_info = [
        "Name of Beneficiary's Settlement",
        "Number of Wells",
        "Total Number of Household Benefitted",
    ]
    table_well_info = doc.add_table(
        rows=len(wells_info) + 1, cols=len(headers_well_info)
    )
    table_well_info.style = "Table Grid"

    hdr_cells = table_well_info.rows[0].cells
    for i, header in enumerate(headers_well_info):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for i, (settlement, num_wells, num_households) in enumerate(wells_info, start=1):
        row_cells = table_well_info.rows[i].cells
        row_cells[0].text = to_utf8(settlement)
        row_cells[1].text = str(num_wells)
        row_cells[2].text = str(num_households)

    doc.add_heading("Detailed Well Information and their Maintenance Demands", level=3)

    all_wells_with_mws_sorted = sorted(
        all_wells_with_mws,
        key=lambda x: (
            not x[0].beneficiary_settlement or x[0].beneficiary_settlement == "NA",
            (x[0].beneficiary_settlement or "").lower(),
        ),
    )

    for i, (well, mws_id) in enumerate(all_wells_with_mws_sorted, 1):
        doc.add_heading(to_utf8(well.beneficiary_settlement), level=4)

        table_well = doc.add_table(rows=15, cols=2)
        table_well.style = "Table Grid"

        def add_well_data(row_idx, label, value):
            row_cells = table_well.rows[row_idx].cells
            row_cells[0].paragraphs[0].add_run(label).bold = True
            row_cells[1].text = to_utf8(value) if value is not None else "NA"

        well_usage = "NA"
        if well.data_well and "Well_usage" in well.data_well:
            well_usage_data = well.data_well["Well_usage"]
            select_one_well_used = ensure_str(
                well_usage_data.get("select_one_well_used")
            )
            select_one_well_used_other = well_usage_data.get(
                "select_one_well_used_other"
            )

            if (
                select_one_well_used
                and select_one_well_used.lower() == "other"
                and select_one_well_used_other
            ):
                well_usage = f"Other: {select_one_well_used_other}"
            elif select_one_well_used:
                well_usage = select_one_well_used

        repair_activities = "NA"
        if well.data_well and "Well_usage" in well.data_well:
            well_usage_data = well.data_well["Well_usage"]
            well_repairs_type = ensure_str(well_usage_data.get("repairs_type"))
            well_repairs_type_other = well_usage_data.get("repairs_type_other")

            if (
                well_repairs_type
                and well_repairs_type.lower() == "other"
                and well_repairs_type_other
            ):
                repair_activities = f"Other: {well_repairs_type_other}"
            elif well_repairs_type:
                repair_activities = well_repairs_type.replace("_", " ")

        # check in the Well_condition
        if (
            repair_activities == "NA"
            and well.data_well
            and "Well_condition" in well.data_well
        ):
            well_condition_data = well.data_well["Well_condition"]
            well_repairs_type = ensure_str(
                well_condition_data.get("select_one_repairs_well")
            )
            well_repairs_type_other = well_condition_data.get(
                "select_one_repairs_well_other"
            )

            if (
                well_repairs_type
                and well_repairs_type.lower() == "other"
                and well_repairs_type_other
            ):
                repair_activities = f"Other: {well_repairs_type_other}"
            elif well_repairs_type:
                repair_activities = well_repairs_type.replace("_", " ")

        add_well_data(0, "MWS ID", mws_id)
        add_well_data(1, "Name of Beneficiary Settlement", well.beneficiary_settlement)
        add_well_data(
            2, "Type of Well", well.data_well.get("select_one_well_type") or "NA"
        )
        add_well_data(3, "Who owns the Well", well.owner)
        add_well_data(
            4, "Beneficiary Name", well.data_well.get("Beneficiary_name") or "NA"
        )
        add_well_data(
            5, "Beneficiary's Father's Name", well.data_well.get("ben_father") or "NA"
        )
        add_well_data(
            6, "Water Availability", well.data_well.get("select_one_year") or "NA"
        )
        add_well_data(7, "Households Benefitted", well.households_benefitted)
        add_well_data(8, "Which Caste uses the well?", well.caste_uses)
        add_well_data(9, "Well Usage", well_usage)
        add_well_data(10, "Need Maintenance?", well.need_maintenance)
        add_well_data(11, "Repair Activities", repair_activities)
        add_well_data(12, "Latitude", well.latitude)
        add_well_data(13, "Longitude", well.longitude)


def populate_consolidated_waterbody_tables(doc, all_waterbodies_with_mws):
    """Create consolidated waterbody summary and detail tables for all MWS IDs"""
    if not all_waterbodies_with_mws:
        doc.add_heading("Water Structure Information", level=3)
        doc.add_paragraph("No water structures found in any of the MWS areas.")
        return

    waterbody_count = defaultdict(int)
    households_count = defaultdict(int)

    for waterbody, mws_id in all_waterbodies_with_mws:
        structure_type = waterbody.water_structure_type
        waterbody_count[(waterbody.beneficiary_settlement, structure_type)] += 1
        households_count[(waterbody.beneficiary_settlement, structure_type)] += int(
            waterbody.household_benefitted
        )

    waterbody_info = [
        (
            settlement,
            waterbody_type,
            waterbody_count[(settlement, waterbody_type)],
            households_count[(settlement, waterbody_type)],
        )
        for (settlement, waterbody_type) in sorted(
            waterbody_count.keys(), key=lambda x: sort_key(x[0])
        )
    ]

    doc.add_heading("Water Structures Summary Information", level=3)
    headers_waterbody_summary = [
        "Name of Beneficiary's Settlement",
        "Type of Water Structure",
        "Number of Waterbodies",
        "Total Number of Household Benefitted",
    ]
    table_waterbody_summary = doc.add_table(
        rows=len(waterbody_info) + 1, cols=len(headers_waterbody_summary)
    )
    table_waterbody_summary.style = "Table Grid"

    hdr_cells = table_waterbody_summary.rows[0].cells
    for i, header in enumerate(headers_waterbody_summary):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for i, (settlement, waterbody_type, num_waterbodies, num_households) in enumerate(
        waterbody_info, start=1
    ):
        row_cells = table_waterbody_summary.rows[i].cells
        row_cells[0].text = to_utf8(settlement)
        row_cells[1].text = to_utf8(waterbody_type)
        row_cells[2].text = str(num_waterbodies)
        row_cells[3].text = str(num_households)

    doc.add_heading("Detailed Water Structures Information", level=3)

    all_waterbodies_with_mws_sorted = sorted(
        all_waterbodies_with_mws, key=lambda x: sort_key(x[0].beneficiary_settlement)
    )

    for i, (waterbody, mws_id) in enumerate(all_waterbodies_with_mws_sorted, 1):
        doc.add_heading(to_utf8(waterbody.beneficiary_settlement), level=4)

        table_water_structure = doc.add_table(rows=14, cols=2)
        table_water_structure.style = "Table Grid"

        def add_waterbody_data(row_idx, label, value):
            row_cells = table_water_structure.rows[row_idx].cells
            row_cells[0].paragraphs[0].add_run(label).bold = True
            row_cells[1].text = to_utf8(value) if value is not None else "NA"

        who_manages = waterbody.who_manages or "NA"
        if who_manages.lower() == "other":
            who_manages = "Other: " + (waterbody.specify_other_manager or "")

        water_structure_type = waterbody.water_structure_type or "NA"
        if water_structure_type.lower() == "other":
            water_structure_type = "Other: " + (waterbody.water_structure_other or "")

        repair_activities = get_waterbody_repair_activities(
            waterbody.data_waterbody, water_structure_type
        )

        add_waterbody_data(0, "MWS ID", mws_id)
        add_waterbody_data(
            1, "Name of the Beneficiary's Settlement", waterbody.beneficiary_settlement
        )
        add_waterbody_data(2, "Who owns the water structure?", waterbody.owner)
        add_waterbody_data(
            3,
            "Beneficiary Name",
            waterbody.data_waterbody.get("Beneficiary_name") or "NA",
        )
        add_waterbody_data(
            4,
            "Beneficiary's Father's Name",
            waterbody.data_waterbody.get("ben_father") or "NA",
        )
        add_waterbody_data(5, "Who manages?", who_manages)
        add_waterbody_data(
            6, "Which Caste uses the water structure?", waterbody.caste_who_uses
        )
        add_waterbody_data(7, "Households Benefitted", waterbody.household_benefitted)
        add_waterbody_data(8, "Type of Water Structure", water_structure_type)
        add_waterbody_data(
            9,
            "Usage of Water Structure",
            format_text(waterbody.data_waterbody.get("select_multiple_uses_structure"))
            or "NA",
        )
        add_waterbody_data(10, "Need Maintenance?", waterbody.need_maintenance)
        add_waterbody_data(11, "Repair Activities", repair_activities)
        add_waterbody_data(12, "Latitude", waterbody.latitude)
        add_waterbody_data(13, "Longitude", waterbody.longitude)


# MARK: - Section E
def add_section_e(doc, plan):
    populate_maintenance_from_waterbody(plan)

    doc.add_heading(
        "Section E: Proposed Maintenance Work",
        level=1,
    )
    para = doc.add_paragraph()
    para.add_run(
        "This section presents information on proposed maintenance works for existing structures based on inputs from the Gram Sabha."
    )
    para.add_run("\n\n")

    asset_types = [
        "Water Recharge Structures",
        "Irrigation Structures",
        "Surface Water Structures",
        "Remote Sensed Surface Water Structures",
    ]

    doc.add_heading("Maintenance Works by Asset Type", level=2)

    table = doc.add_table(rows=1 + len(asset_types), cols=1)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Asset Type"
    header_cells[0].paragraphs[0].runs[0].bold = True

    for i, asset_type in enumerate(asset_types):
        table.rows[i + 1].cells[0].text = asset_type

    for asset_type in asset_types:
        doc.add_heading(f"Maintenance Works for {asset_type}", level=3)

        if asset_type == "Water Recharge Structures":
            maintenance_gw_table(doc, plan)
        elif asset_type == "Irrigation Structures":
            maintenance_agri_table(doc, plan)
        elif asset_type == "Surface Water Structures":
            maintenance_waterstructures_table(doc, plan)
        elif asset_type == "Remote Sensed Surface Water Structures":
            maintenance_rs_waterstructures_table(doc, plan)

        doc.add_page_break()


def maintenance_gw_table(doc, plan):
    headers = [
        "Type of demand",
        "Name of the Beneficiary Settlement",
        "Beneficiary Name",
        "Beneficiary's Father's Name",
        "Type of Recharge Structure",
        "Repair Activities",
        "Latitude",
        "Longitude",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    for m in get_maintenance_data(plan.id, "gw"):
        row_cells = table.add_row().cells
        row_cells[0].text = to_utf8(m["demand_type"] or "NA")
        row_cells[1].text = to_utf8(m["beneficiary_settlement"] or "NA")
        row_cells[2].text = to_utf8(m["beneficiary_name"] or "NA")
        row_cells[3].text = to_utf8(m["beneficiary_father_name"] or "NA")
        row_cells[4].text = to_utf8(m["structure_type"] or "NA")
        row_cells[5].text = to_utf8(m["repair_activities"] or "NA")
        row_cells[6].text = str(m["latitude"])
        row_cells[7].text = str(m["longitude"])


def maintenance_agri_table(doc, plan):
    headers = [
        "Type of demand",
        "Name of the Beneficiary Settlement",
        "Beneficiary Name",
        "Beneficiary's Father's Name",
        "Type of Irrigation Structure",
        "Repair Activity",
        "Latitude",
        "Longitude",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    for m in get_maintenance_data(plan.id, "agri"):
        row_cells = table.add_row().cells
        row_cells[0].text = to_utf8(m["demand_type"] or "NA")
        row_cells[1].text = to_utf8(m["beneficiary_settlement"] or "NA")
        row_cells[2].text = to_utf8(m["beneficiary_name"] or "NA")
        row_cells[3].text = to_utf8(m["beneficiary_father_name"] or "NA")
        row_cells[4].text = to_utf8(m["structure_type"] or "NA")
        row_cells[5].text = to_utf8(m["repair_activities"] or "NA")
        row_cells[6].text = str(m["latitude"])
        row_cells[7].text = str(m["longitude"])


def maintenance_waterstructures_table(doc, plan):
    headers = [
        "Type of demand",
        "Name of the Beneficiary Settlement",
        "Beneficiary Name",
        "Beneficiary's Father's Name",
        "Type of Work",
        "Repair Activities",
        "Latitude",
        "Longitude",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    for m in get_maintenance_data(plan.id, "swb"):
        row_cells = table.add_row().cells
        row_cells[0].text = to_utf8(m["demand_type"] or "NA")
        row_cells[1].text = to_utf8(m["beneficiary_settlement"] or "NA")
        row_cells[2].text = to_utf8(m["beneficiary_name"] or "NA")
        row_cells[3].text = to_utf8(m["beneficiary_father_name"] or "NA")
        row_cells[4].text = to_utf8(m["structure_type"] or "NA")
        row_cells[5].text = to_utf8(m["repair_activities"] or "NA")
        row_cells[6].text = str(m["latitude"])
        row_cells[7].text = str(m["longitude"])


def maintenance_rs_waterstructures_table(doc, plan):
    headers = [
        "Type of demand",
        "Name of the Beneficiary Settlement",
        "Beneficiary Name",
        "Beneficiary's Father's Name",
        "Type of Work",
        "Repair Activities",
        "Latitude",
        "Longitude",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    for m in get_maintenance_data(plan.id, "swb_rs"):
        row_cells = table.add_row().cells
        row_cells[0].text = to_utf8(m["demand_type"] or "NA")
        row_cells[1].text = to_utf8(m["beneficiary_settlement"] or "NA")
        row_cells[2].text = to_utf8(m["beneficiary_name"] or "NA")
        row_cells[3].text = to_utf8(m["beneficiary_father_name"] or "NA")
        row_cells[4].text = to_utf8(m["structure_type"] or "NA")
        row_cells[5].text = to_utf8(m["repair_activities"] or "NA")
        row_cells[6].text = str(m["latitude"])
        row_cells[7].text = str(m["longitude"])


# MARK: - Section F
def add_section_f(doc, plan, mws):
    doc.add_heading("Section F: Proposed New NRM works on basis through Gram Sabha")
    para = doc.add_paragraph()
    para.add_run(
        "This section outlines the details of proposed new NRM works based on community inputs put up to the Gram Sabha.\n\n"
    )

    create_nrm_works_table(doc, plan, mws)


def create_nrm_works_table(doc, plan, mws):
    works = get_nrm_works_data(plan.id)

    headers = [
        "S.No",
        "Work Category : Irrigation work or Recharge Structure",
        "Type of demand",
        "Work demand",
        "Name of Beneficiary's Settlement",
        "Beneficiary's Name",
        "Gender",
        "Beneficiary's Father's Name",
        "Latitude",
        "Longitude",
    ]
    table = doc.add_table(rows=1 + len(works), cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    for i, w in enumerate(works, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[1].text = to_utf8(w["work_category"])
        row_cells[2].text = to_utf8(w["demand_type"] or "NA")
        row_cells[3].text = to_utf8(w["work_demand"] or "NA")
        row_cells[4].text = to_utf8(w["beneficiary_settlement"] or "NA")
        row_cells[5].text = to_utf8(w["beneficiary_name"] or "NA")
        row_cells[6].text = to_utf8(w["gender"] or "NA")
        row_cells[7].text = to_utf8(w["beneficiary_father_name"] or "NA")
        row_cells[8].text = str(w["latitude"])
        row_cells[9].text = str(w["longitude"])


# MARK: - Section G -- Plantations and Livelihood Works
def add_section_g(doc, plan, mws):
    doc.add_heading("Section G: Proposed New Livelihood Works", level=1)

    all_livelihood = get_livelihood_data(plan.id)
    livestock_fisheries = [
        r for r in all_livelihood if r["livelihood_work"] in ("Livestock", "Fisheries")
    ]
    plantations_etc = [
        r
        for r in all_livelihood
        if r["livelihood_work"] not in ("Livestock", "Fisheries")
    ]

    doc.add_heading("G.1 Livestock and Fisheries", level=2)
    lf_headers = [
        "Livelihood Works",
        "Type of Demand",
        "Work Demand",
        "Name of Beneficiary Settlement",
        "Beneficiary's Name",
        "Gender",
        "Beneficiary Father's Name",
        "Latitude",
        "Longitude",
    ]
    table = doc.add_table(rows=1, cols=len(lf_headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(lf_headers):
        hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for r in livestock_fisheries:
        row_cells = table.add_row().cells
        row_cells[0].text = to_utf8(r["livelihood_work"])
        row_cells[1].text = to_utf8(r["demand_type"] or "NA")
        row_cells[2].text = to_utf8(r["work_demand"] or "NA")
        row_cells[3].text = to_utf8(r["beneficiary_settlement"] or "NA")
        row_cells[4].text = to_utf8(r["beneficiary_name"] or "NA")
        row_cells[5].text = to_utf8(r["gender"] or "NA")
        row_cells[6].text = to_utf8(r["beneficiary_father_name"] or "NA")
        row_cells[7].text = "{:.6f}".format(r["latitude"]) if r["latitude"] else "NA"
        row_cells[8].text = "{:.6f}".format(r["longitude"]) if r["longitude"] else "NA"

    doc.add_heading("G.2 Plantations and Kitchen Gardens", level=2)
    plantation_headers = [
        "Livelihood Works",
        "Type of demand",
        "Name of Beneficiary Settlement",
        "Name of Beneficiary",
        "Gender",
        "Beneficiary's Father's Name",
        "Name of Plantation Crop",
        "Total Acres",
        "Latitude",
        "Longitude",
    ]
    plantation_table = doc.add_table(rows=1, cols=len(plantation_headers))
    plantation_table.style = "Table Grid"
    plantation_hdr_cells = plantation_table.rows[0].cells
    for i, header in enumerate(plantation_headers):
        plantation_hdr_cells[i].paragraphs[0].add_run(header).bold = True

    for r in plantations_etc:
        row_cells = plantation_table.add_row().cells
        row_cells[0].text = to_utf8(r["livelihood_work"])
        row_cells[1].text = to_utf8(r["demand_type"] or "NA")
        row_cells[2].text = to_utf8(r["beneficiary_settlement"] or "NA")
        row_cells[3].text = to_utf8(r["beneficiary_name"] or "NA")
        row_cells[4].text = to_utf8(r["gender"] or "NA")
        row_cells[5].text = to_utf8(r["beneficiary_father_name"] or "NA")
        row_cells[6].text = to_utf8(r["work_demand"] or "NA")
        row_cells[7].text = to_utf8(r["total_acres"] or "NA")
        row_cells[8].text = "{:.6f}".format(r["latitude"]) if r["latitude"] else "NA"
        row_cells[9].text = "{:.6f}".format(r["longitude"]) if r["longitude"] else "NA"
